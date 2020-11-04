from docx import Document
from duckpy import Client

client = Client()

startstr = "1"
endstr = "2"
sconv = int(startstr)
econv = int(endstr)

print("\nCopyright © 2020 gabcode. All rights reserved.")
print("The unauthorized copying, sharing or distribution of copyrighted material is strictly prohibited.\n")
print("IMPORTANT: ONLY WORD (.docx) DOCUMENTS ARE SUPPORTED FOR NOW \n")
print("INSTRUCTIONS: ")
print("1. Copy and paste your .docx file into the same folder of this program.")
print("2. Make sure to write the name of the document correctly or the program will close")
print("3. All the questions must have the corresponding number before the question \n")
filename = input("What is the document name? (don't write .docx at the end): ") + ".docx"
the_type = input("What type of document is this? (questions / definitions): ")
subject = input("Tell me the subject of this document: ")

if the_type.lower() == "questions":
    num_of_questions = input("What is the number of questions?: ")

elif the_type.lower() == "definitions":
    num_of_questions = input("What is the number of definitions?: ")

try:
    doc = Document(filename)
except ValueError:
    print("Wrong file name. Restarting...")
    input()

    import sys

    sys.exit(0)

docFinal = Document()


def ReadingTextDocuments():
    completedText = []

    for paragraph in doc.paragraphs:
        completedText.append(paragraph.text)
    return '\n'.join(completedText)


s = ReadingTextDocuments()
# Question recognizer code
# try:

for i in range(int(num_of_questions)):
    start = s.find(str(sconv))
    end = s.find(str(econv))
    substring = s[start:end]
    sconv += 1
    econv += 1
    print(substring)

    # to search
    if the_type.lower() == "definitions":
        results = client.search(f'define ∼{substring} +{subject}')

        print(results[0]["description"])

        paragraph = docFinal.add_paragraph(results[0]["description"])
        prior_paragraph = paragraph.insert_paragraph_before(substring)

    elif the_type.lower() == "questions":
        results = client.search(f'∼{substring}')

        print(results[0]["description"])

        paragraph = docFinal.add_paragraph(results[0]["description"])
        prior_paragraph = paragraph.insert_paragraph_before(substring)



# except ValueError:
# print("Couldn't recognize questions. Restarting...")
# input()

# import sys

# sys.exit(0)

docFinal.save("completed-" + filename)
input('\n Press ENTER to exit')
